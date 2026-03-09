from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)


def build_manual(output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual de Usuario\nObservatorio de Precios (No Tecnico)")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(24, 55, 109)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Version operativa - {datetime.now().strftime('%d/%m/%Y')}\n"
        "Uso diario del scraper para iPhone 17 y Samsung S25"
    )
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(68, 84, 106)

    doc.add_paragraph("")

    callout_table = doc.add_table(rows=1, cols=1)
    callout_cell = callout_table.cell(0, 0)
    _set_cell_shading(callout_cell, "E8F0FE")
    callout_text = callout_cell.paragraphs[0].add_run(
        "Resumen rapido: este proceso ya esta preparado para uso diario. "
        "Con un solo comando se generan CSV, JSON y HTML con fecha/hora."
    )
    callout_text.bold = True
    callout_text.font.color.rgb = RGBColor(24, 55, 109)

    doc.add_heading("1. Para que sirve", level=1)
    doc.add_paragraph(
        "Este sistema sirve para comparar precios de una oferta concreta de productos "
        "entre Santander Boutique y competidores."
    )
    doc.add_paragraph(
        "Oferta actual monitorizada (scope fijo):", style="List Bullet"
    )
    doc.add_paragraph("Apple iPhone 17 en todas sus variantes (incluye iPhone Air).", style="List Bullet")
    doc.add_paragraph("Samsung Galaxy S25 en todas sus variantes disponibles.", style="List Bullet")

    doc.add_heading("2. Que hace automaticamente en cada ejecucion", level=1)
    doc.add_paragraph("Cada vez que se lanza el script:")
    doc.add_paragraph("1) Extrae productos base desde Santander Boutique.", style="List Number")
    doc.add_paragraph("2) Busca esos mismos productos en competidores.", style="List Number")
    doc.add_paragraph("3) Genera archivos de salida para analisis e historico.", style="List Number")
    doc.add_paragraph("4) Actualiza un HTML de comparativa listo para compartir.", style="List Number")

    doc.add_heading("3. Instalacion inicial (solo una vez)", level=1)
    doc.add_paragraph("Pide a IT o a una persona tecnica que ejecute estos comandos una vez:")
    _add_code_block(
        doc,
        "python -m pip install -r requirements.full.txt\n"
        "python -m playwright install chromium",
    )

    doc.add_heading("4. Ejecucion diaria (paso a paso)", level=1)
    doc.add_paragraph("Paso 1. Abre una terminal en la carpeta del proyecto.")
    doc.add_paragraph("Paso 2. Ejecuta este comando:")
    _add_code_block(doc, "python run_observatorio_focus_fast.py")
    doc.add_paragraph("Paso 3. Espera a ver lineas [OK] al final de la ejecucion.")

    doc.add_heading("5. Archivos que se generan", level=1)
    doc.add_paragraph("El sistema deja dos tipos de salida:")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    table.rows[0].cells[0].text = "Tipo de archivo"
    table.rows[0].cells[1].text = "Ruta"
    rows = [
        ("Ultima ejecucion (JSON)", "output/latest_prices.json"),
        ("Ultima ejecucion (CSV)", "output/latest_prices.csv"),
        ("Historico JSON", "output/prices_YYYYMMDD_HHMMSS.json"),
        ("Historico CSV", "output/prices_YYYYMMDD_HHMMSS.csv"),
        ("HTML ultima ejecucion", "output/price_comparison_live.html"),
        ("HTML historico", "output/price_comparison_live_YYYYMMDD_HHMMSS.html"),
    ]
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right

    doc.add_heading("6. Checklist diaria (2 minutos)", level=1)
    doc.add_paragraph("Revisa esto al terminar:")
    doc.add_paragraph("[ ] Aparece [OK] Registros de precio en consola.", style="List Bullet")
    doc.add_paragraph("[ ] Se ha creado CSV/JSON historico con fecha y hora.", style="List Bullet")
    doc.add_paragraph("[ ] Se ha creado HTML historico con fecha y hora.", style="List Bullet")
    doc.add_paragraph("[ ] El HTML abre y muestra productos de la tirada.", style="List Bullet")

    doc.add_heading("7. Como compartir resultados", level=1)
    doc.add_paragraph("Para enviar la corrida del dia, comparte estos 2 archivos:")
    doc.add_paragraph("1) CSV historico de esa ejecucion.", style="List Number")
    doc.add_paragraph("2) HTML historico de esa ejecucion.", style="List Number")
    doc.add_paragraph(
        "Ejemplo: output/prices_20260225_111420.csv y output/price_comparison_live_20260225_111420.html"
    )

    doc.add_heading("8. Si algo falla (guia simple)", level=1)
    doc.add_paragraph("Caso A. No arranca el navegador (Playwright):")
    _add_code_block(doc, "python -m playwright install chromium")
    doc.add_paragraph("Caso B. Un competidor aparece con 0 resultados:")
    doc.add_paragraph(
        "Esto puede pasar por bloqueos anti-bot o cambios de web. Repite la tirada. "
        "Si persiste, avisa al equipo tecnico."
    )
    doc.add_paragraph("Caso C. Va mas lento de lo esperado:")
    doc.add_paragraph("Puedes ejecutar por marca para dividir carga:", style="List Bullet")
    _add_code_block(
        doc,
        "python run_observatorio_focus_fast.py --brand Samsung\n"
        "python run_observatorio_focus_fast.py --brand Apple",
    )

    doc.add_heading("9. Como mejorarlo sin romper lo que funciona", level=1)
    doc.add_paragraph("Regla clave: primero validar en pequeno, luego escalar.")
    doc.add_paragraph("Orden recomendado:", style="List Bullet")
    doc.add_paragraph("1) Prueba en un solo competidor.", style="List Number")
    doc.add_paragraph("2) Prueba en una sola marca.", style="List Number")
    doc.add_paragraph("3) Prueba completa.", style="List Number")
    doc.add_paragraph(
        "Nunca cambies a la vez scraping + plantilla HTML + estructura de salida en una sola iteracion."
    )

    doc.add_heading("10. Comandos utiles", level=1)
    _add_code_block(
        doc,
        "python run_observatorio_focus_fast.py\n"
        "python run_observatorio_focus_fast.py --brand Samsung\n"
        "python run_observatorio_focus_fast.py --brand Apple\n"
        "python run_observatorio_focus_fast.py --competitors \"Santander Boutique,Amazon,Media Markt\"",
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Observatorio de Precios - Manual de usuario no tecnico")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    out = Path("output/doc/Manual_Usuario_Observatorio_NoTecnico.docx")
    path = build_manual(out)
    print(f"[OK] DOCX generado: {path}")
